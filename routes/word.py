from flask import Blueprint, request, jsonify, send_file, current_app, make_response
from docx import Document
from io import BytesIO, StringIO
from utils.auth import get_user_from_token
from pymongo import MongoClient
import traceback
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml import OxmlElement
from lxml import etree
import csv
import pandas as pd
from datetime import datetime, timezone
import pytz
from bson.objectid import ObjectId
import os
import re
import zipfile
from lxml import etree
import shutil
import tempfile

# Blueprint for Word document generation routes
bp = Blueprint('word', __name__, url_prefix='/word')

def get_db():
    """
    Central function to get database connection consistently
    throughout the application
    """
    client = MongoClient(current_app.config["MONGODB_URI"])
    db = client['sushanto']  # Use consistent database name
    return db, client

def get_template_path(template_type, region=None):
    """
    Get the appropriate template file path based on template type and region.
    
    Args:
        template_type (str): Type of template ('Global', 'Regional', or 'Country')
        region (str, optional): Region name for Regional templates
        
    Returns:
        str: Path to the template file
    
    Raises:
        ValueError: If template type is invalid or if region is invalid
    """
    if template_type == 'Regional':
        if not region:
            raise ValueError("Region is required for Regional template")
            
        # Map regions to their specific template files
        region_template_mapping = {
            'North America': 'north_america_region_template.docx',
            'Europe': 'europe_region_template.docx',
            'Asia Pacific': 'asia_pacific_region_template.docx',
            'Middle East & Africa': 'middle_east_africa_region_template.docx',
            'Latin America': 'latin_america_region_template.docx'
        }
        
        if region not in region_template_mapping:
            raise ValueError(f"Invalid region: {region}")
            
        template_filename = region_template_mapping[region]
    else:
        # Handle Global and Country templates
        template_mapping = {
            'Global': 'global_template.docx',
            'Country': 'country_template.docx'
        }
        
        if template_type not in template_mapping:
            raise ValueError(f"Invalid template type: {template_type}")
            
        template_filename = template_mapping[template_type]
    
    template_path = os.path.join(current_app.root_path, 'templates', template_filename)
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file not found: {template_filename}")
        
    return template_path

def validate_region(region):
    """
    Validate and normalize region name.
    
    Args:
        region (str): Region name to validate
        
    Returns:
        str: Normalized region name
        
    Raises:
        ValueError: If region is invalid
    """
    valid_regions = {
        "North America",
        "Europe",
        "Asia Pacific",
        "Middle East & Africa",
        "Latin America"
    }
    
    if region not in valid_regions:
        raise ValueError(f"Invalid region: {region}. Must be one of {valid_regions}")
    
    return region

def replace_region(doc, region):
    """Replace region placeholder in the document."""
    current_app.logger.info(f"Replacing region placeholder with: {region}")
    replace_text(doc, "{{region}}", region)

def replace_country(doc, country):
    """Replace country placeholder in the document."""
    current_app.logger.info(f"Replacing country placeholder with: {country}")
    replace_text(doc, "{{country}}", country)

def replace_textbox_text(doc, placeholder, replacement):
    """
    Replace placeholders in textboxes throughout the Word document including headers.
    """
    current_app.logger.info(f"Starting replacement process for '{placeholder}'")
    replacement_done = False
    
    namespaces = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
        'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
    }

    def process_element(element):
        nonlocal replacement_done
        textbox_paths = [
            './/w:txbxContent',
            './/mc:AlternateContent//w:txbxContent',
            './/wps:txbx//w:txbxContent'
        ]
        
        for path in textbox_paths:
            for txbox in element.findall(path, namespaces):
                current_app.logger.debug(f"Found textbox with path: {path}")
                for t in txbox.findall('.//w:t', namespaces):
                    if t.text and placeholder in t.text:
                        current_app.logger.info(f"Found placeholder in text: {t.text}")
                        t.text = t.text.replace(placeholder, replacement)
                        replacement_done = True
                        current_app.logger.info(f"Replaced with: {t.text}")

    # Process headers in all sections
    for section in doc.sections:
        if section.header:
            try:
                # Process header part
                header_element = section.header._element
                process_element(header_element)
                
                # Also process any header references
                if hasattr(section.header, 'part'):
                    process_element(section.header.part.element)
            except Exception as e:
                current_app.logger.error(f"Error processing header: {str(e)}", exc_info=True)

    # Process main document
    for part in doc.part.package.parts:
        try:
            if hasattr(part, 'element'):
                process_element(part.element)
        except Exception as e:
            current_app.logger.error(f"Error processing part {part.partname}: {str(e)}", exc_info=True)
    
    if not replacement_done:
        current_app.logger.warning("No replacements were made")
    
    return replacement_done

def replace_text(doc, placeholder, replacement):
    """Replace placeholders throughout the document and log replacements."""
    replacement_done = False

    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            current_app.logger.debug(f"Replacing '{placeholder}' with '{replacement}' in paragraph: {paragraph.text}")
            paragraph.text = paragraph.text.replace(placeholder, replacement)
            replacement_done = True

    for section in doc.sections:
        header = section.header
        if header:
            for paragraph in header.paragraphs:
                if placeholder in paragraph.text:
                    current_app.logger.debug(f"Replacing '{placeholder}' with '{replacement}' in header paragraph: {paragraph.text}")
                    paragraph.text = paragraph.text.replace(placeholder, replacement)
                    replacement_done = True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        current_app.logger.debug(f"Replacing '{placeholder}' with '{replacement}' in table cell: {paragraph.text}")
                        paragraph.text = paragraph.text.replace(placeholder, replacement)
                        replacement_done = True

    if not replacement_done:
        current_app.logger.warning(f"No occurrences of '{placeholder}' found in the document for replacement.")

def replace_segments(doc, segments):
    """Replace placeholders for segments and sub-segments in the document."""
    for i, segment in enumerate(segments):
        segment_key = f"{{{{Segment{i + 1}}}}}"
        segment_name = segment.get("name", "")
        current_app.logger.info(f"Replacing segment placeholder: {segment_key} with '{segment_name}'")
        replace_text(doc, segment_key, segment_name)

        sub_segments = segment.get("subSegments", [])
        for j, subsegment in enumerate(sub_segments):
            sub_key = f"{{{{Segment{i + 1}Sub-segment{j + 1}}}}}"
            current_app.logger.info(f"Replacing sub-segment placeholder: {sub_key} with '{subsegment}'")
            replace_text(doc, sub_key, subsegment)

def replace_companies(doc, companies):
    """Replace placeholders for company names in the document."""
    for i, company in enumerate(companies[:10]):  # Limit to 10 companies
        company_key = f"{{{{Company{i + 1}}}}}"
        current_app.logger.info(f"Replacing company placeholder: {company_key} with '{company}'")
        replace_text(doc, company_key, company)

def clean_empty_segments(doc, user_inputs):
    """
    Removes paragraphs and table rows containing unprocessed placeholders 
    for segments and sub-segments.
    """
    segment_placeholders = []

    # Generate all possible placeholders for segments 1-6 and sub-segments 1-10
    for i in range(1, 7):  # Segments 1-6
        segment_placeholders.append(f"{{{{Segment{i}}}}}") 
        # Also add variation with lowercase 's'
        segment_placeholders.append(f"{{{{Segment{i}}}}}") 
        for j in range(1, 11):  # Sub-Segments 1-10
            # Add both variants of sub-segment formatting
            segment_placeholders.append(f"{{{{Segment{i}Sub-Segment{j}}}}}")
            segment_placeholders.append(f"{{{{Segment{i}Sub-segment{j}}}}}")

    current_app.logger.info(f"Checking for these placeholders: {segment_placeholders}")

    # Remove paragraphs containing unprocessed placeholders
    paragraphs_to_remove = []  
    for para in doc.paragraphs:
        para_text = para.text.strip()
        if not para_text:
            continue
        
        current_app.logger.info(f"Processing paragraph: '{para_text}'")

        # Check if any placeholder is still in the text
        found_placeholders = [ph for ph in segment_placeholders if ph in para_text]
        if found_placeholders:
            current_app.logger.info(f"Found unprocessed placeholders in paragraph: {found_placeholders}")
            current_app.logger.info(f"Marking paragraph for removal: '{para_text}'")
            paragraphs_to_remove.append(para)

    # Remove marked paragraphs
    for para in paragraphs_to_remove:
        p = para._element
        p.getparent().remove(p)
        current_app.logger.info(f"Removed paragraph: '{para.text.strip()}'")

    # Remove table rows containing unprocessed placeholders
    for table in doc.tables:
        rows_to_remove = []  
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(row_cells)
            if not row_text.strip():
                continue
            
            current_app.logger.info(f"Processing table row: '{row_text}'")

            # Check each cell individually
            found_placeholders = []
            for cell in row.cells:
                cell_text = " ".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                for placeholder in segment_placeholders:
                    if placeholder in cell_text:
                        found_placeholders.append(placeholder)
                        current_app.logger.info(f"Found placeholder in cell: {placeholder}")

            if found_placeholders:
                current_app.logger.info(f"Found unprocessed placeholders in row: {found_placeholders}")
                current_app.logger.info(f"Marking row for removal: '{row_text}'")
                rows_to_remove.append(row)

        # Remove marked rows
        for row in rows_to_remove:
            tr = row._element
            tr.getparent().remove(tr)
            current_app.logger.info(f"Removed table row with placeholders")

    current_app.logger.info("Finished cleaning empty segments and table rows.")

def remove_unused_sections(doc, user_inputs):
    """
    Removes entire sections of content between {{SegmentX_Start}} and {{SegmentX_End}} 
    markers for segments that aren't present in user_inputs.
    
    Args:
        doc: The Word document object
        user_inputs: List of segment dictionaries containing user input data
    """
    # Get the list of segment numbers that are present in user inputs
    present_segments = set()
    for i, segment in enumerate(user_inputs):
        if segment.get("name"):  # If segment has content
            present_segments.add(i + 1)
    
    current_app.logger.info(f"Present segments: {present_segments}")
    
    # Find and store all paragraph and table elements between markers
    for segment_num in range(1, 7):  # For segments 1-6
        if segment_num not in present_segments:
            start_marker = f"{{{{Segment{segment_num}_Start}}}}"
            end_marker = f"{{{{Segment{segment_num}_End}}}}"
            
            # Process each segment that needs to be removed
            all_elements = list(doc.element.body)
            
            # Track markers - we may have multiple instances of the same segment
            segment_zones = []
            start_idx = None
            
            # First pass: Identify all zones to remove
            for idx, element in enumerate(all_elements):
                # Check if element is a paragraph
                if element.tag.endswith('p'):
                    paragraph_idx = len([e for e in all_elements[:idx] if e.tag.endswith('p')])
                    if paragraph_idx < len(doc.paragraphs):
                        paragraph = doc.paragraphs[paragraph_idx]
                        
                        if start_marker in paragraph.text and start_idx is None:
                            start_idx = idx
                        
                        if end_marker in paragraph.text and start_idx is not None:
                            segment_zones.append((start_idx, idx))
                            start_idx = None  # Reset for potential next instance
            
            # Second pass: Remove the identified zones in reverse order
            # (removing from end to start to avoid index shifting issues)
            current_app.logger.info(f"Found {len(segment_zones)} instances of Segment {segment_num} to remove")
            
            for start_idx, end_idx in sorted(segment_zones, reverse=True):
                elements_to_remove = all_elements[start_idx:end_idx+1]
                
                current_app.logger.info(f"Removing section for Segment {segment_num} (indices {start_idx}-{end_idx})")
                
                for element_to_remove in elements_to_remove:
                    if element_to_remove.getparent() is not None:
                        element_to_remove.getparent().remove(element_to_remove)

def clean_all_segment_markers(doc):
    """
    Removes all segment markers (start and end) from the document,
    including those for present segments.
    """
    for segment_num in range(1, 7):
        start_marker = f"{{{{Segment{segment_num}_Start}}}}"
        end_marker = f"{{{{Segment{segment_num}_End}}}}"
        
        # Remove from paragraphs
        for paragraph in doc.paragraphs:
            if start_marker in paragraph.text or end_marker in paragraph.text:
                p = paragraph._element
                if p.getparent() is not None:
                    p.getparent().remove(p)
                    current_app.logger.info(f"Removed marker: {paragraph.text}")
        
        # Remove from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if start_marker in paragraph.text:
                            paragraph.text = paragraph.text.replace(start_marker, '')
                        if end_marker in paragraph.text:
                            paragraph.text = paragraph.text.replace(end_marker, '')

def update_document_references(doc):
    """
    Updates Table of Contents, List of Figures, and List of Tables in a Word document.
    Uses Word field codes to force an update of these elements.
    
    Args:
        doc: The Word document object
    """
    try:
        # XML namespace for Word documents
        namespace = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        }
        
        # Function to update specific field type
        def update_fields(field_type):
            # Find all fields of the specified type
            elements = doc.element.body.findall(
                f'.//w:fldChar/..',  # Parent of fldChar
                namespace
            )
            
            current_app.logger.info(f"Found {len(elements)} potential field elements")
            
            for element in elements:
                # Get the field code text
                field_codes = element.findall('.//w:instrText', namespace)
                if not field_codes:
                    continue
                    
                field_text = ' '.join(code.text for code in field_codes if code.text)
                
                # Check if this is the type of field we want to update
                if field_type in field_text:
                    current_app.logger.info(f"Updating {field_type}")
                    
                    # Find the fldChar elements
                    fld_chars = element.findall('.//w:fldChar', namespace)
                    
                    # Update the field status
                    for fld_char in fld_chars:
                        fld_char_type = fld_char.get(f'{{{namespace["w"]}}}fldCharType', '')
                        if fld_char_type == 'begin':
                            # Set dirty attribute to trigger update
                            fld_char.set(f'{{{namespace["w"]}}}dirty', 'true')
                            current_app.logger.info(f"Marked {field_type} field as dirty for update")
        
        # Update Table of Contents
        update_fields('TOC')
        
        # Update List of Figures
        update_fields('TOF')
        
        # Update List of Tables
        update_fields('TOT')
        
        # Additional update for any other potential field codes
        for field_type in ['REF', 'PAGEREF', 'SEQ']:
            update_fields(field_type)
            
        current_app.logger.info("Document references update complete")
        
    except Exception as e:
        current_app.logger.error(f"Error updating document references: {str(e)}")
        raise

def store_document_data(db, user_id, input_data, filename, template_type, generation_type="single"):
    """
    Store document generation data in MongoDB.
    
    Args:
        db: MongoDB database instance
        user_id: ID of the user generating the document
        input_data: Dictionary containing the input data used to generate the document
        filename: Name of the generated document
        template_type: Type of template used (Global, Regional, Country)
        generation_type: Type of generation (single or bulk)
    """
    document_record = {
        "user_id": ObjectId(user_id),
        "input_data": input_data,
        "filename": filename,
        "template_type": template_type,
        "generation_type": generation_type,
        "created_at": datetime.now(pytz.timezone('Asia/Kolkata')).isoformat(),  # IST timezone
        "status": "completed"
    }
    
    try:
        result = db.documents.insert_one(document_record)
        return str(result.inserted_id)
    except Exception as e:
        current_app.logger.error(f"Error storing document data: {str(e)}")
        raise

def replace_header_textbox(doc, market_name, location):
    """
    Replace placeholders {{region}}/{{country}} and {{market_name}} in textboxes located 
    in headers, particularly those positioned above images.
    
    Parameters:
        doc (str): Path to the Word document
        market_name (str): The market name to replace {{market_name}} placeholder
        location (str): The region or country name to replace {{region}} or {{country}} placeholder
    
    Returns:
        str: Path to the modified document
    """
    import os
    import re
    import zipfile
    from lxml import etree
    import shutil
    
    # Import Flask current_app if available, otherwise create a fallback logger
    try:
        from flask import current_app
        logger = current_app.logger
    except ImportError:
        import logging
        logging.basicConfig(
            level=logging.DEBUG,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        logger = logging.getLogger(__name__)
    
    logger.info(f"Starting header textbox replacement: market_name={market_name}, location={location}")
    
    # Build the replacements dictionary
    replacements = {}
    if market_name:
        replacements['{{market_name}}'] = market_name
    
    if location:
        # Include both region and country placeholders
        replacements['{{region}}'] = location
        replacements['{{country}}'] = location
    
    # Exit early if no replacements needed
    if not replacements:
        logger.warning("No replacement values provided")
        return doc
    
    # Handle different input types
    is_path = isinstance(doc, str)
    doc_path = None
    
    if is_path:
        # It's a file path
        doc_path = doc
        base, ext = os.path.splitext(doc_path)
        output_path = f"{base}_filled{ext}"
        # Create a copy of the document to work with
        shutil.copy2(doc_path, output_path)
        working_path = output_path
    else:
        # It's a Document object, save to temp file
        import tempfile
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        working_path = temp_file.name
        temp_file.close()
        doc.save(working_path)
    
    # Add alternate forms of the placeholders
    expanded_replacements = replacements.copy()
    for placeholder, value in replacements.items():
        # Remove curly braces for alternate form
        alt_key = placeholder.replace('{{', '').replace('}}', '')
        # Create alternate forms
        expanded_replacements[f"{{{{{alt_key}}}}}"] = value
        expanded_replacements[f"{{ {{{alt_key}}} }}"] = value
        expanded_replacements[f"{{{{{alt_key.replace('_', '')}}}}}"] = value  # No underscores
        expanded_replacements[alt_key] = value  # Plain text version
        
        # Variant with no underscore and alternative casing
        no_underscore = alt_key.replace('_', '')
        expanded_replacements[f"{{{{{no_underscore}}}}}"] = value
        expanded_replacements[f"{{{{{no_underscore.lower()}}}}}" ] = value
        
        logger.debug(f"Expanded placeholders for {placeholder}: {expanded_replacements}")
    
    # Word document XML namespaces
    namespaces = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'v': 'urn:schemas-microsoft-com:vml',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
        'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
        'wpg': 'http://schemas.microsoft.com/office/word/2010/wordprocessingGroup',
        'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
        'a14': 'http://schemas.microsoft.com/office/drawing/2010/main',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
        'o': 'urn:schemas-microsoft-com:office:office',
        'w10': 'urn:schemas-microsoft-com:office:word',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    }
    
    # Helper function to extract paragraph text
    def extract_paragraph_text(paragraph):
        text = ""
        for run in paragraph.xpath('.//w:r', namespaces=namespaces):
            for text_elem in run.xpath('.//w:t', namespaces=namespaces):
                if text_elem.text:
                    text += text_elem.text
        return text
    
    # Helper function to check if paragraph contains any placeholder
    def contains_any_placeholder(text, placeholders):
        # Direct matching
        for placeholder in placeholders:
            if placeholder in text:
                return True
        
        # Check for split placeholders (e.g., "{{reg" "ion}}")
        for placeholder in placeholders:
            if placeholder.startswith('{{') and '{{' in text:
                placeholder_core = placeholder.strip('{}')
                # Check if any substantial part of the placeholder exists
                if len(placeholder_core) > 3 and placeholder_core[:3] in text:
                    return True
        
        return False
    
    # Process the document
    with zipfile.ZipFile(output_path, mode='a') as zip_file:
        # Process headers
        for i in range(1, 4):  # Usually Word docs have up to 3 headers
            header_path = f'word/header{i}.xml'
            
            try:
                # Skip if header doesn't exist
                if header_path not in zip_file.namelist():
                    continue
                
                logger.info(f"Processing header: {header_path}")
                
                # Extract the XML content
                xml_content = zip_file.read(header_path)
                root = etree.fromstring(xml_content)
                
                header_modified = False
                
                # Process different types of textboxes
                textbox_paths = [
                    # VML textboxes often used in older Word docs
                    ('.//v:textbox//w:p', "VML textbox"),
                    # DrawingML textboxes 
                    ('.//w:txbxContent//w:p', "DrawingML textbox"),
                    # Shapes with textboxes
                    ('.//wps:wsp//wps:txbx//w:txbxContent//w:p', "shape textbox"),
                    # Alternate content (compatibility mode)
                    ('.//mc:AlternateContent//w:txbxContent//w:p', "alternate content textbox")
                ]
                
                # Process each type of textbox
                for xpath, context in textbox_paths:
                    for paragraph in root.xpath(xpath, namespaces=namespaces):
                        paragraph_text = extract_paragraph_text(paragraph)
                        
                        # Skip if no placeholders found
                        if not contains_any_placeholder(paragraph_text, expanded_replacements.keys()):
                            continue
                        
                        logger.info(f"Found placeholder in {context}: {paragraph_text}")
                        
                        # Replace all placeholders in the text
                        processed_text = paragraph_text
                        for placeholder, replacement in expanded_replacements.items():
                            if placeholder in processed_text:
                                processed_text = processed_text.replace(placeholder, replacement)
                                logger.debug(f"Replaced '{placeholder}' with '{replacement}'")
                        
                        # If no direct replacements were made, try for split placeholders
                        if processed_text == paragraph_text:
                            logger.debug(f"Attempting to handle split placeholders in {context}")
                            
                            for placeholder, replacement in expanded_replacements.items():
                                if placeholder.startswith('{{') and '{{' in paragraph_text:
                                    placeholder_core = placeholder.strip('{}')
                                    # Look for fragments of the placeholder
                                    for i in range(3, len(placeholder_core), 3):
                                        fragment = placeholder_core[:i]
                                        if fragment in paragraph_text:
                                            start_pos = paragraph_text.find(fragment)
                                            # Check if this is likely the start of our placeholder
                                            prefix = paragraph_text[max(0, start_pos-2):start_pos]
                                            if '{{' in prefix or start_pos < 2:
                                                logger.debug(f"Found fragment '{fragment}' of placeholder '{placeholder}'")
                                                # Replace the placeholder and surrounding braces
                                                before = max(0, paragraph_text.rfind('{{', 0, start_pos+1))
                                                after = paragraph_text.find('}}', start_pos)
                                                if after > -1:
                                                    processed_text = (paragraph_text[:before] + 
                                                                      replacement + 
                                                                      paragraph_text[after+2:])
                                                    break
                        
                        # Skip if no changes needed
                        if processed_text == paragraph_text:
                            continue
                            
                        logger.info(f"Replacing with: {processed_text}")
                        
                        # Update the text content
                        runs = paragraph.xpath('.//w:r', namespaces=namespaces)
                        if runs:
                            # Put all content in first run and clear others
                            first_run = runs[0]
                            text_elements = first_run.xpath('.//w:t', namespaces=namespaces)
                            
                            if text_elements:
                                text_elements[0].text = processed_text
                                # Clear other runs to avoid duplicated content
                                for run in runs[1:]:
                                    for text_elem in run.xpath('.//w:t', namespaces=namespaces):
                                        text_elem.text = ""
                                
                                header_modified = True
                                logger.info(f"Updated content in {context}")
                
                # Handle DrawingML text elements (often used for text over images)
                drawing_text_paths = [
                    ('.//a:p//a:r//a:t', 'DrawingML text'),
                    ('.//wp:inline//a:p//a:r//a:t', 'Inline drawing text'),
                    ('.//wp:anchor//a:p//a:r//a:t', 'Anchored drawing text')
                ]
                
                for xpath, context in drawing_text_paths:
                    for text_elem in root.xpath(xpath, namespaces=namespaces):
                        if not text_elem.text:
                            continue
                            
                        original_text = text_elem.text
                        modified_text = original_text
                        
                        # Check for direct replacements
                        placeholder_found = False
                        for placeholder, replacement in expanded_replacements.items():
                            if placeholder in modified_text:
                                modified_text = modified_text.replace(placeholder, replacement)
                                placeholder_found = True
                                logger.info(f"Found placeholder {placeholder} in {context}")
                        
                        # Update single text element if modified
                        if original_text != modified_text:
                            text_elem.text = modified_text
                            header_modified = True
                            logger.info(f"Updated text in {context} from '{original_text}' to '{modified_text}'")
                
                # If changes were made, update the header in the zip file
                if header_modified:
                    updated_xml = etree.tostring(root, encoding='UTF-8', xml_declaration=True)
                    zip_file.writestr(header_path, updated_xml)
                    logger.info(f"Updated content in header {header_path}")
            
            except Exception as e:
                logger.error(f"Error processing header {header_path}: {str(e)}", exc_info=True)
    
    logger.info(f"Document processing completed, output saved to {output_path}")
    return output_path



def log_placeholders(doc):
    """Log all placeholders found in the document for debugging."""
    for paragraph in doc.paragraphs:
        current_app.logger.info(f"Paragraph text: {paragraph.text}")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    current_app.logger.info(f"Table cell text: {paragraph.text}")

@bp.route('/generate', methods=['POST'])
def generate_word_doc():

    if request.method == 'OPTIONS':
        response = make_response()
        response.headers.add('Access-Control-Allow-Origin', '*')  # Or your specific frontend origin
        response.headers.add('Access-Control-Allow-Methods', 'POST, OPTIONS')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type, Authorization')
        response.headers.add('Access-Control-Max-Age', '3600')
        return response
    
    """Generate a Word document based on template and input data."""
    db, client = get_db()
    documents_collection = db.documents
    token = request.headers.get('Authorization')

    if not token:
        return jsonify({"message": "Token is required"}), 401

    user = get_user_from_token(token)
    if not user:
        return jsonify({"message": "Invalid token"}), 401

    data = request.get_json()
    current_app.logger.info(f"Input data: {data}")
    if not data:
        current_app.logger.warning(f"User {user['username']} attempted to create a document without any data")
        return jsonify({"message": "No data provided"}), 400
    
    # Get template type from request data
    template_type = data.get("template_type", "Global")  # Default to Global if not specified

    try:
        # Get the appropriate template path
        template_path = get_template_path(
            template_type,
            region=data.get("region") if template_type == "Regional" else None
        )
        
        # Load the template document
        doc = Document(template_path)
        log_placeholders(doc)

        

        # Handle region-specific content for Regional template
        if template_type == "Regional":
            region = data.get("region")
            if not region:
                raise ValueError("Region is required for Regional template")
            region = validate_region(region)
            replace_region(doc, region)

                       

        # Handle country-specific content for Country template
        if template_type == "Country":
            country = data.get("country")
            if not country:
                raise ValueError("Country is required for Country template")
            replace_country(doc, country)
            
            
            
        # Extract market_name
        market_name = data.get("market_name", "")
        if market_name:
            replace_text(doc, "{{market_name}}", market_name)
            replace_textbox_text(doc, "{{market_name}}", market_name)
            template_type = data.get("template_type", "Global")
           
                  
            
        # Convert flat dictionary to structured segmentations
        segmentations = []
        i = 1
        while f"Segment{i}" in data:
            segment_name = data[f"Segment{i}"]
            sub_segments = []
            j = 1
            while f"Segment{i}Sub-segment{j}" in data:
                sub_segments.append(data[f"Segment{i}Sub-segment{j}"])
                j += 1
            segmentations.append({"name": segment_name, "subSegments": sub_segments})
            i += 1

        current_app.logger.info(f"Converted segmentations: {segmentations}")

        # Replace segments in the document
        replace_segments(doc, segmentations)

        # Clean empty paragraphs and table rows
        clean_empty_segments(doc, segmentations)
        
        # Replace company names
        companies = [data.get(f"Company{i+1}", "") for i in range(10)]
        replace_companies(doc, companies)

        # Remove unused sections and clean markers
        remove_unused_sections(doc, segmentations)
        clean_all_segment_markers(doc)

        update_document_references(doc)


       # Create the temporary file first
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_path = temp_file.name
        temp_file.close()
        doc.save(temp_path)

        # Now handle the textbox replacements based on template type
        market_name = data.get("market_name", "")
        updated_path = None

        if template_type == "Country":
            country = data.get("country", "")
            if country:  # Only proceed if country exists
                updated_path = replace_header_textbox(temp_path, market_name, country)
        elif template_type == "Global":
            global_sample = "Global"
            updated_path = replace_header_textbox(temp_path, market_name, global_sample)
        elif template_type == "Regional":
            region = data.get("region", "")
            if region:  # Only proceed if region exists
                updated_path = replace_header_textbox(temp_path, market_name, region)

        # Process the updated document if successful
        if updated_path and os.path.exists(updated_path):
            doc = Document(updated_path)
            # Clean up temp files
            os.remove(temp_path)
            os.remove(updated_path)
        else:
            # If no successful update, use the original temp document
            doc = Document(temp_path)
        
        # Save the updated document to an in-memory file
        in_memory_file = BytesIO()
        doc.save(in_memory_file)
        in_memory_file.seek(0)

        current_app.logger.info(f"Document generated successfully by user {user['username']} using {template_type} template")
        current_app.logger.info(f"Market Name: {market_name}")
        
       # Generate filename based on template type
        if template_type == "Regional":
            filename = f'{data.get("region")} {market_name} Market.docx'
        elif template_type == "Country":
            filename = f'{data.get("country")} {market_name} Market.docx'
        else:  # Global template
            filename = f'Global {market_name} Market.docx'

        # Store document data in MongoDB
        doc_id = store_document_data(
            db=db,
            user_id=user['_id'],
            input_data=data,
            filename=filename,
            template_type=data.get("template_type", "Global")
        )
            
        return send_file(
            in_memory_file,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            download_name=filename,
            as_attachment=True
        )
    
        # Add CORS headers to the send_file response
        origin = request.headers.get('Origin')
        if origin in current_app.config.get('ALLOWED_ORIGINS', []):
            current_app.logger.info(f"Adding CORS headers for file response to origin: {origin}")
            response.headers.set('Access-Control-Allow-Origin', origin)
            response.headers.set('Access-Control-Allow-Credentials', 'false')
        
        return response
    
    except FileNotFoundError as e:
        current_app.logger.error(f"Template file not found: {str(e)}")
        return jsonify({'message': 'Template file not found', 'error': str(e)}), 404
    except ValueError as e:
        current_app.logger.error(f"Invalid template type or region: {str(e)}")
        return jsonify({'message': 'Invalid template type or region', 'error': str(e)}), 400
    
 
    
    except Exception as e:
        current_app.logger.error(f"Error generating document by user {user['username']}: {traceback.format_exc()}")
        return jsonify({'message': 'Error generating document', 'error': str(e)}), 500
        # Update document status in case of failure
        if 'doc_id' in locals():
            db.documents.update_one(
                {"_id": ObjectId(doc_id)},
                {"$set": {"status": "failed", "error": str(e)}}
            )
        return jsonify({'message': 'Error generating document', 'error': str(e)}), 500
    
    # Add CORS headers to error response
    origin = request.headers.get('Origin')
    if origin in current_app.config.get('ALLOWED_ORIGINS', []):
        response.headers.set('Access-Control-Allow-Origin', origin)
        response.headers.set('Access-Control-Allow-Credentials', 'false')
    
    return response, 500
    
@bp.route('/generate-bulk', methods=['POST'])
def generate_bulk_documents():
    """Generate multiple Word documents from CSV data."""
    try:
        db, client = get_db()
        documents_collection = db.documents
        
        # Add logging for all headers
        current_app.logger.info(f"All headers: {dict(request.headers)}")
        
        # Get and validate token
        token = request.headers.get('Authorization')
        current_app.logger.info(f"Token received: {token[:10]}..." if token else "No token")
        
        if not token:
            return jsonify({"message": "Token is required"}), 401
        
        # Log before user validation
        current_app.logger.info("About to validate user from token")
        user = get_user_from_token(token)
        current_app.logger.info(f"User validation result: {user is not None}")
        
        if not user:
            return jsonify({"message": "Invalid token"}), 401

        current_app.logger.info(f"Processing bulk document generation for user: {user['username']}")

        if 'file' not in request.files:
            return jsonify({"message": "No file provided"}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({"message": "No file selected"}), 400

        # Try different encodings
        encodings_to_try = ['utf-8', 'utf-8-sig', 'latin1', 'cp1252']
        csv_data = None
        successful_encoding = None

        for encoding in encodings_to_try:
            try:
                # Reset file pointer
                file.stream.seek(0)
                # Try to read with current encoding
                stream = StringIO(file.stream.read().decode(encoding))
                csv_data = pd.read_csv(stream)
                successful_encoding = encoding
                break
            except UnicodeDecodeError:
                continue
            except Exception as e:
                current_app.logger.error(f"Error reading CSV with {encoding} encoding: {str(e)}")
                continue

        if csv_data is None:
            return jsonify({
                "message": "Unable to read CSV file. Please ensure it's properly encoded (UTF-8 or Windows-1252)."
            }), 400

        current_app.logger.info(f"Successfully read CSV with {successful_encoding} encoding")

        # Validate required columns
        required_columns = ['template_type', 'market_name']
        missing_columns = [col for col in required_columns if col not in csv_data.columns]
        if missing_columns:
            return jsonify({
                "message": f"Missing required columns: {', '.join(missing_columns)}"
            }), 400

        # Create a ZIP file in memory
        memory_zip = BytesIO()
        with zipfile.ZipFile(memory_zip, 'w') as zf:
            # Track processing results
            results = {
                'success': 0,
                'failed': 0,
                'errors': []
            }
           
            # Create a bulk operation record
            bulk_record = {
                "user_id": ObjectId(user['_id']),
                "type": "bulk",
                "status": "processing",
                "created_at": datetime.now(pytz.timezone('Asia/Kolkata')).isoformat(),  # IST timezone
                "total_files": 0,
                "successful_files": 0,
                "failed_files": 0,
                "files": []
            }
        
            bulk_id = db.documents.insert_one(bulk_record).inserted_id    

            # Process each row in the CSV
            for index, row in csv_data.iterrows():
                try:
                    # Define filename early to use in error handling
                    filename = "undefined_file.docx"  # Default filename
                    
                    # Validate template type
                    if row['template_type'] not in ['Global', 'Regional', 'Country']:
                        raise ValueError(f"Invalid template_type in row {index + 1}: {row['template_type']}")

                    # Validate region for Regional template
                    if row['template_type'] == 'Regional' and (pd.isna(row.get('region')) or row['region'].strip() == ''):
                        raise ValueError(f"Region is required for Regional template in row {index + 1}")

                    # Validate country for Country template
                    if row['template_type'] == 'Country' and (pd.isna(row.get('country')) or row['country'].strip() == ''):
                        raise ValueError(f"Country is required for Country template in row {index + 1}")

                    # Prepare data for document generation
                    doc_data = {
                        "template_type": row['template_type'],
                        "market_name": row['market_name'],
                        "region": row.get('region', ''),
                        "country": row.get('country', '')
                    }
                    
                    # Add segments and sub-segments
                    for i in range(1, 7):
                        segment_key = f'segment{i}'
                        if segment_key in row and pd.notna(row[segment_key]):
                            doc_data[f'Segment{i}'] = str(row[segment_key]).strip()
                            # Add sub-segments if they exist
                            for j in range(1, 11):
                                sub_key = f'segment{i}_sub{j}'
                                if sub_key in row and pd.notna(row[sub_key]):
                                    doc_data[f'Segment{i}Sub-segment{j}'] = str(row[sub_key]).strip()

                    # Add companies
                    for i in range(1, 11):
                        company_key = f'company{i}'
                        if company_key in row and pd.notna(row[company_key]):
                            doc_data[f'Company{i}'] = str(row[company_key]).strip()

                    # Create filename BEFORE database operations
                    if doc_data['template_type'] == 'Regional':
                        filename = f"{doc_data['region']}_{doc_data['market_name']}_Market.docx"
                    elif doc_data['template_type'] == 'Country':
                        filename = f"{doc_data['country']}_{doc_data['market_name']}_Market.docx"
                    else:
                        filename = f"{doc_data['market_name']}_Global_Market.docx"

                    # Clean filename
                    filename = "".join(c for c in filename if c.isalnum() or c in (' ', '_', '-', '.'))
                    
                    # Store individual document data
                    doc_record = {
                        "user_id": ObjectId(user['_id']),
                        "bulk_id": bulk_id,
                        "input_data": row.to_dict(),
                        "filename": filename,
                        "template_type": row['template_type'],
                        "generation_type": "bulk",
                        "created_at": datetime.now(pytz.timezone('Asia/Kolkata')).isoformat(),  # IST timezone
                        "status": "completed"
                    }
                
                    doc_id = db.documents.insert_one(doc_record).inserted_id
                    
                    # Generate document
                    doc = generate_single_document(doc_data, user)
                    
                    # Add to ZIP
                    zf.writestr(filename, doc.getvalue())
                    results['success'] += 1
                    
                    # Update bulk record for successful file
                    db.documents.update_one(
                        {"_id": bulk_id},
                        {
                            "$inc": {
                                "total_files": 1,
                                "successful_files": 1
                            },
                            "$push": {
                                "files": {
                                    "doc_id": doc_id,
                                    "filename": filename,
                                    "status": "completed"
                                }
                            }
                        }
                    )

                except Exception as e:
                    error_msg = f"Error in row {index + 1}: {str(e)}"
                    current_app.logger.error(error_msg)
                    results['failed'] += 1
                    results['errors'].append(error_msg)
                    
                    # Update bulk record for failed file
                    db.documents.update_one(
                        {"_id": bulk_id},
                        {
                            "$inc": {
                                "total_files": 1,
                                "failed_files": 1
                            },
                            "$push": {
                                "files": {
                                    "filename": filename,  # Now safely defined
                                    "status": "failed",
                                    "error": str(e)
                                }
                            }
                        }
                    )

            # Update final status of bulk operation
            db.documents.update_one(
                {"_id": bulk_id},
                {"$set": {"status": "completed"}}
            )

            # If all documents failed, return error
            if results['failed'] > 0 and results['success'] == 0:
                return jsonify({
                    "message": "Failed to generate any documents",
                    "errors": results['errors']
                }), 500

        # Prepare ZIP file for download
        memory_zip.seek(0)
        
        # Log completion
        current_app.logger.info(
            f"Bulk generation completed for user: {user['username']}. "
            f"Success: {results['success']}, Failed: {results['failed']}"
        )
        
        return send_file(
            memory_zip,
            mimetype='application/zip',
            as_attachment=True,
            download_name='generated_documents.zip'
        )

    except Exception as e:
        current_app.logger.error(f"Error in bulk generation: {str(e)}")
        if 'bulk_id' in locals():
            db.documents.update_one(
                {"_id": bulk_id},
                {"$set": {"status": "failed", "error": str(e)}}
            )
        return jsonify({
            "message": "Error processing bulk generation request",
            "error": str(e)
        }), 500

def generate_single_document(data, user):
    """Generate a single document and return it as BytesIO object."""
    try:
        template_path = get_template_path(
            data["template_type"],
            region=data.get("region") if data["template_type"] == "Regional" else None
        )
        
        doc = Document(template_path)
        
        # Extract all placeholder values upfront
        region = data.get("region")
        country = data.get("country")
        market_name = data.get("market_name", "")
        
        # Process region/country specific content
        if data["template_type"] == "Regional":
            region = validate_region(region)
            replace_region(doc, region)
            # Handle textbox replacements using dictionary format
            replace_header_textbox(doc, {"region": region})
            
           
                
        elif data["template_type"] == "Country":
            replace_country(doc, country)
            # Handle textbox replacements using dictionary format
            replace_header_textbox(doc, {"country": country})
            
            
        # Replace market name
        if market_name:
            replace_text(doc, "{{market_name}}", market_name)
            # Use dictionary form for consistency
            replace_header_textbox(doc, {"market_name": market_name})
            replace_textbox_text(doc, "{{market_name}}", market_name)

        # Process segments and companies
        segmentations = []
        for i in range(1, 7):
            if f'Segment{i}' in data:
                sub_segments = []
                for j in range(1, 11):
                    sub_key = f'Segment{i}Sub-segment{j}'
                    if sub_key in data:
                        sub_segments.append(data[sub_key])
                segmentations.append({
                    "name": data[f'Segment{i}'],
                    "subSegments": sub_segments
                })

        replace_segments(doc, segmentations)
        clean_empty_segments(doc, segmentations)
        
        companies = [data.get(f"Company{i}", "") for i in range(1, 11)]
        replace_companies(doc, companies)

        remove_unused_sections(doc, segmentations)
        clean_all_segment_markers(doc)
        update_document_references(doc)

        # Save to BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    except Exception as e:
        current_app.logger.error(f"Error generating single document: {str(e)}")
        raise